#!/usr/bin/env python3
"""
DOCX Content Extractor

Extracts content from DOCX files maintaining exact document order.
All elements (paragraphs, images, tables) are stored in a single
ordered list to preserve their positions as they appear in the document.
"""

from __future__ import annotations

import io
import os
import re
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from zipfile import ZipFile

# python-docx for DOCX parsing
try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.table import Table, _Cell
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. Install with: pip install python-docx")

# PIL for image processing
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# lxml for XML manipulation
from lxml import etree


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class ExtractedImage:
    """Represents an extracted image from DOCX."""
    filename: str
    data: bytes
    content_type: str
    width: Optional[int] = None
    height: Optional[int] = None
    alt_text: str = ""
    caption: str = ""
    rel_id: str = ""  # relationship ID in DOCX


@dataclass
class ExtractedTable:
    """Represents an extracted table from DOCX."""
    rows: List[List[str]]
    header_rows: int = 1
    caption: str = ""
    has_merged_cells: bool = False
    num_cols: int = 0


@dataclass
class TextBlock:
    """Represents a block of text with formatting."""
    text: str
    style: str = "Normal"
    level: int = 0  # Heading level (0 = not a heading)
    is_bold: bool = False
    is_italic: bool = False
    is_underline: bool = False
    list_type: Optional[str] = None  # "bullet" or "number" or None
    list_level: int = 0
    alignment: str = "left"


@dataclass
class DocumentElement:
    """
    A single element in document order.
    Can be a paragraph, image, or table.
    Only one of (paragraph, image, table) will be set.
    """
    element_type: str  # "paragraph", "image", "table"
    paragraph: Optional[TextBlock] = None
    image: Optional[ExtractedImage] = None
    table: Optional[ExtractedTable] = None


@dataclass
class DocxContent:
    """Complete extracted content from a DOCX file."""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    # ALL elements in document order (paragraphs, images, tables mixed)
    elements: List[DocumentElement] = field(default_factory=list)

    # Separate lists for convenience (also populated)
    images: List[ExtractedImage] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    text_blocks: List[TextBlock] = field(default_factory=list)

    # Structure tracking
    chapters: List[Dict[str, Any]] = field(default_factory=list)


# ============================================================================
# DOCX EXTRACTOR CLASS
# ============================================================================

class DocxExtractor:
    """
    Extracts content from DOCX files maintaining exact document order.

    All elements are processed in the order they appear in the DOCX body,
    preserving the relationship between text, images, and tables.
    """

    def __init__(
        self,
        extract_images: bool = True,
        extract_tables: bool = True,
        preserve_formatting: bool = True,
        min_image_size: int = 50
    ):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.preserve_formatting = preserve_formatting
        self.min_image_size = min_image_size

        self._image_counter = 0
        self._table_counter = 0

    def extract(self, docx_path: Union[str, Path]) -> DocxContent:
        """
        Extract all content from a DOCX file in document order.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            DocxContent with all extracted content in order
        """
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        # Reset counters
        self._image_counter = 0
        self._table_counter = 0

        # Open document
        doc = Document(str(docx_path))

        content = DocxContent()

        # Extract metadata
        content.metadata = self._extract_metadata(doc)
        content.title = content.metadata.get("title", docx_path.stem)
        content.authors = content.metadata.get("authors", [])

        # Extract images from the DOCX ZIP package
        image_data_map = {}  # rel_id -> ExtractedImage
        if self.extract_images:
            image_data_map = self._extract_images_from_package(docx_path, doc)

        # Process document body in order
        self._process_body_in_order(doc, content, image_data_map)

        return content

    def _extract_metadata(self, doc: DocumentType) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {}

        try:
            core_props = doc.core_properties

            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["authors"] = [core_props.author]
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
        except Exception as e:
            print(f"Warning: Could not extract all metadata: {e}")

        return metadata

    def _extract_images_from_package(self, docx_path: Path, doc: DocumentType) -> Dict[str, ExtractedImage]:
        """
        Extract all images from the DOCX package and map them to relationship IDs.

        Returns a dict mapping relationship IDs to ExtractedImage objects.
        """
        # First, extract raw image data from ZIP
        media_files = {}  # partname -> (data, content_type, ext)
        try:
            with ZipFile(docx_path, 'r') as zf:
                for name in zf.namelist():
                    if name.startswith('word/media/'):
                        img_data = zf.read(name)
                        ext = Path(name).suffix.lower()
                        content_types = {
                            '.png': 'image/png',
                            '.jpg': 'image/jpeg',
                            '.jpeg': 'image/jpeg',
                            '.gif': 'image/gif',
                            '.bmp': 'image/bmp',
                            '.tiff': 'image/tiff',
                            '.tif': 'image/tiff',
                            '.emf': 'image/x-emf',
                            '.wmf': 'image/x-wmf',
                        }
                        content_type = content_types.get(ext, 'application/octet-stream')

                        # Get image dimensions
                        width, height = None, None
                        if HAS_PIL and ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                            try:
                                img = Image.open(io.BytesIO(img_data))
                                width, height = img.size
                                if width < self.min_image_size and height < self.min_image_size:
                                    continue
                            except Exception:
                                pass

                        media_files[name] = (img_data, content_type, ext, width, height)
        except Exception as e:
            print(f"Warning: Error reading DOCX package: {e}")

        # Build a quick lookup by filename from media_files
        media_by_name = {}
        for media_path, file_data in media_files.items():
            basename = Path(media_path).name
            media_by_name[basename] = (media_path, file_data)

        # Map relationship IDs to media files
        rel_to_image = {}
        try:
            rels = doc.part.rels
            for rel_id, rel in rels.items():
                # Check relationship type - only process image relationships
                rel_type = getattr(rel, 'reltype', '') or ''
                is_image_rel = 'image' in rel_type.lower()

                if hasattr(rel, 'target_part') and rel.target_part:
                    target = rel.target_part
                    if hasattr(target, 'partname'):
                        partname = str(target.partname)
                        target_name = Path(partname).name

                        # Match against media files by filename
                        if target_name in media_by_name:
                            media_path, (data, ct, ext, w, h) = media_by_name[target_name]
                            self._image_counter += 1
                            img = ExtractedImage(
                                filename=f"img_{self._image_counter:04d}{ext}",
                                data=data,
                                content_type=ct,
                                width=w,
                                height=h,
                                rel_id=rel_id
                            )
                            rel_to_image[rel_id] = img
                        elif is_image_rel:
                            # Try partial match for renamed/relocated images
                            for media_path, (data, ct, ext, w, h) in media_files.items():
                                if media_path.endswith(target_name):
                                    self._image_counter += 1
                                    img = ExtractedImage(
                                        filename=f"img_{self._image_counter:04d}{ext}",
                                        data=data,
                                        content_type=ct,
                                        width=w,
                                        height=h,
                                        rel_id=rel_id
                                    )
                                    rel_to_image[rel_id] = img
                                    break
                elif is_image_rel and hasattr(rel, 'target_ref'):
                    # External image link - try to match by name
                    target_ref = str(rel.target_ref)
                    target_name = Path(target_ref).name
                    if target_name in media_by_name:
                        media_path, (data, ct, ext, w, h) = media_by_name[target_name]
                        self._image_counter += 1
                        img = ExtractedImage(
                            filename=f"img_{self._image_counter:04d}{ext}",
                            data=data,
                            content_type=ct,
                            width=w,
                            height=h,
                            rel_id=rel_id
                        )
                        rel_to_image[rel_id] = img
        except Exception as e:
            print(f"Warning: Could not map relationship IDs: {e}")

        print(f"  - Mapped {len(rel_to_image)} image relationships from {len(media_files)} media files")
        return rel_to_image

    def _process_body_in_order(self, doc: DocumentType, content: DocxContent, image_data_map: Dict[str, ExtractedImage]):
        """
        Process the document body in exact element order.
        Paragraphs, images, and tables are all added to the elements list
        in the order they appear in the DOCX.
        """
        body = doc.element.body

        # Build lookup dictionaries for O(1) access
        para_lookup = {para._element: para for para in doc.paragraphs}
        table_lookup = {table._element: table for table in doc.tables}

        element_count = 0
        total_elements = len(list(body))
        print(f"  - Processing {total_elements} body elements...")

        for element in body:
            element_count += 1
            if element_count % 500 == 0:
                print(f"  - Processing element {element_count}/{total_elements}...")

            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # Check for images in this paragraph FIRST
                if self.extract_images:
                    found_images = self._find_images_in_element(element, image_data_map)
                    for img in found_images:
                        content.images.append(img)
                        content.elements.append(DocumentElement(
                            element_type="image",
                            image=img
                        ))

                # Extract paragraph text
                para = para_lookup.get(element)
                if para is None:
                    try:
                        para = Paragraph(element, doc)
                    except Exception:
                        continue

                block = self._extract_paragraph(para)
                if block and (block.text.strip() or block.list_type):
                    content.text_blocks.append(block)
                    content.elements.append(DocumentElement(
                        element_type="paragraph",
                        paragraph=block
                    ))

                    # Track chapters
                    if block.level >= 1:
                        self._track_chapter(content, block)

            elif tag == 'tbl' and self.extract_tables:
                table = table_lookup.get(element)
                if table is None:
                    try:
                        table = Table(element, doc)
                    except Exception:
                        continue

                extracted_table = self._extract_table(table)
                if extracted_table and extracted_table.rows:
                    self._table_counter += 1
                    content.tables.append(extracted_table)
                    content.elements.append(DocumentElement(
                        element_type="table",
                        table=extracted_table
                    ))

        print(f"  - Extracted: {len(content.text_blocks)} paragraphs, {len(content.images)} images, {len(content.tables)} tables")

    def _find_images_in_element(self, element, image_data_map: Dict[str, ExtractedImage]) -> List[ExtractedImage]:
        """
        Find all images in a paragraph element using multiple detection methods:
        1. DrawingML: a:blip with r:embed or r:link
        2. VML: v:imagedata with r:id (legacy format)
        3. w:object with embedded images

        Returns list of ExtractedImage objects found.
        """
        found = []
        seen_rel_ids = set()

        # Method 1: DrawingML images (a:blip)
        blips = element.findall('.//' + qn('a:blip'))
        for blip in blips:
            # Check r:embed first (most common)
            rel_id = blip.get(qn('r:embed'))
            if not rel_id:
                # Also check r:link (linked images)
                rel_id = blip.get(qn('r:link'))
            if rel_id and rel_id in image_data_map and rel_id not in seen_rel_ids:
                seen_rel_ids.add(rel_id)
                found.append(image_data_map[rel_id])

        # Method 2: VML images (v:imagedata) - legacy DOCX format
        try:
            # VML namespace: urn:schemas-microsoft-com:vml
            vml_ns = 'urn:schemas-microsoft-com:vml'
            r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

            imagedata_elems = element.findall('.//{%s}imagedata' % vml_ns)
            for imgdata in imagedata_elems:
                rel_id = imgdata.get('{%s}id' % r_ns)
                if not rel_id:
                    rel_id = imgdata.get('{%s}pict' % r_ns)
                if rel_id and rel_id in image_data_map and rel_id not in seen_rel_ids:
                    seen_rel_ids.add(rel_id)
                    found.append(image_data_map[rel_id])
        except Exception:
            pass

        # Method 3: OLE objects with image representations
        try:
            ole_objects = element.findall('.//' + qn('w:object'))
            for obj in ole_objects:
                # OLE objects can contain v:shape with v:imagedata
                vml_shapes = obj.findall('.//{%s}imagedata' % 'urn:schemas-microsoft-com:vml')
                for imgdata in vml_shapes:
                    r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                    rel_id = imgdata.get('{%s}id' % r_ns)
                    if rel_id and rel_id in image_data_map and rel_id not in seen_rel_ids:
                        seen_rel_ids.add(rel_id)
                        found.append(image_data_map[rel_id])
        except Exception:
            pass

        return found

    def _extract_paragraph(self, para: Paragraph) -> Optional[TextBlock]:
        """Extract a paragraph with its formatting."""
        # Get paragraph style
        style_name = para.style.name if para.style else "Normal"

        # Determine heading level
        level = 0
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                if "1" in style_name:
                    level = 1
                elif "2" in style_name:
                    level = 2
                elif "3" in style_name:
                    level = 3
        elif style_name == "Title":
            level = 1
        elif style_name == "Subtitle":
            level = 2

        # Check for list
        list_type = None
        list_level = 0

        numPr = para._element.find('.//' + qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None:
                list_level = int(ilvl.get(qn('w:val'), '0'))
            list_type = "bullet"

            numId = numPr.find(qn('w:numId'))
            if numId is not None:
                try:
                    num_val = int(numId.get(qn('w:val'), '0'))
                    if num_val % 2 == 0 or "List Number" in style_name or "Numbered" in style_name:
                        list_type = "number"
                except ValueError:
                    pass

        if "List Bullet" in style_name or "Bullet" in style_name:
            list_type = "bullet"
        elif "List Number" in style_name or "Numbered" in style_name:
            list_type = "number"

        # Extract text - use para.text as primary (most reliable)
        text = para.text or ""

        # Get formatting info from runs
        is_bold = False
        is_italic = False
        is_underline = False

        if self.preserve_formatting and para.runs:
            text_parts = []
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                if run.bold:
                    is_bold = True
                    run_text = f"**{run_text}**"
                if run.italic:
                    is_italic = True
                    run_text = f"*{run_text}*"
                if run.underline:
                    is_underline = True
                text_parts.append(run_text)

            if text_parts:
                text = "".join(text_parts)
                text = re.sub(r'\*\*\*\*+', '', text)
                text = re.sub(r'\*\*\s*\*\*', '', text)
                text = re.sub(r'\*\s*\*', '', text)
        else:
            for run in para.runs:
                if run.bold:
                    is_bold = True
                if run.italic:
                    is_italic = True
                if run.underline:
                    is_underline = True

        # Get alignment
        alignment = "left"
        if para.alignment:
            if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                alignment = "center"
            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                alignment = "right"
            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                alignment = "justify"

        return TextBlock(
            text=text,
            style=style_name,
            level=level,
            is_bold=is_bold,
            is_italic=is_italic,
            is_underline=is_underline,
            list_type=list_type,
            list_level=list_level,
            alignment=alignment
        )

    def _extract_table(self, table: Table) -> Optional[ExtractedTable]:
        """Extract a table with its structure."""
        rows = []
        has_merged = False

        try:
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_cells.append(cell_text)
                if row_cells:
                    rows.append(row_cells)
        except Exception as e:
            print(f"Warning: Error extracting table: {e}")
            return None

        if not rows:
            return None

        num_cols = max(len(row) for row in rows) if rows else 1

        return ExtractedTable(
            rows=rows,
            header_rows=1,
            has_merged_cells=has_merged,
            num_cols=num_cols
        )

    def _track_chapter(self, content: DocxContent, block: TextBlock):
        """Track chapter/section structure."""
        if block.level == 1:
            content.chapters.append({
                "title": block.text,
                "level": 1,
                "sections": []
            })
        elif block.level >= 2 and content.chapters:
            content.chapters[-1]["sections"].append({
                "title": block.text,
                "level": block.level
            })


# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def extract_docx(docx_path: Union[str, Path], **kwargs) -> DocxContent:
    """
    Convenience function to extract content from a DOCX file.

    Args:
        docx_path: Path to the DOCX file
        **kwargs: Additional arguments passed to DocxExtractor

    Returns:
        DocxContent with all extracted content
    """
    extractor = DocxExtractor(**kwargs)
    return extractor.extract(docx_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <docx_file>")
        sys.exit(1)

    docx_path = sys.argv[1]
    content = extract_docx(docx_path)

    print(f"Title: {content.title}")
    print(f"Authors: {content.authors}")
    print(f"Total elements: {len(content.elements)}")
    print(f"  Paragraphs: {len(content.text_blocks)}")
    print(f"  Images: {len(content.images)}")
    print(f"  Tables: {len(content.tables)}")
    print(f"  Chapters: {len(content.chapters)}")

    # Print first few elements
    print("\nFirst 10 elements:")
    for i, elem in enumerate(content.elements[:10]):
        if elem.element_type == "paragraph":
            prefix = "#" * elem.paragraph.level if elem.paragraph.level > 0 else "P"
            print(f"  {i+1}. [{prefix}] {elem.paragraph.text[:80]}")
        elif elem.element_type == "image":
            print(f"  {i+1}. [IMG] {elem.image.filename}")
        elif elem.element_type == "table":
            print(f"  {i+1}. [TBL] {len(elem.table.rows)} rows x {elem.table.num_cols} cols")
